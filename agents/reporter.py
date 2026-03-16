"""
Report Agent
────────────
Compiles a professional .docx report with:
  - Cover page
  - Executive summary
  - Methodology
  - Key findings
  - Embedded charts (PNG → inline image)
  - Conclusions
  - Appendix (code excerpt)

Falls back gracefully if python-docx is unavailable.
"""

from __future__ import annotations

import json
import logging
import re
import tempfile
from pathlib import Path
from typing import Any

from agents.base import BaseAgent
from agents.state import AgentState

logger = logging.getLogger(__name__)

_SYSTEM_PROMPT = """\
You are a professional technical report writer.
Write a detailed Markdown report based on the provided information.

Use EXACTLY these section headings (## prefix):
## Executive Summary
## Methodology
## Key Findings & Insights
## Data Overview
## Conclusions & Recommendations
## Appendix

Rules:
- Write a FULL report — at least 500 words
- Use the research summary as the primary knowledge source
- Never mention pipeline errors, retries, or technical failures
- Use bullet points inside sections where appropriate
- The Appendix section should contain a brief code description only
"""


class ReporterAgent(BaseAgent):
    LLM_ROLE = "primary"

    def __init__(self) -> None:
        super().__init__("ReporterAgent")

    def run(self, state: AgentState) -> dict[str, Any]:
        task        = state.get("original_task", "") or ""
        subtasks    = state.get("subtasks", []) or []
        research    = state.get("research_summary", "") or ""
        stdout      = state.get("execution_stdout", "") or ""
        artifacts   = state.get("execution_artifacts", []) or []
        code        = state.get("generated_code", "") or ""
        token_usage = state.get("token_usage", {}) or {}
        timings     = state.get("step_timings", {}) or {}
        sources     = state.get("research_sources", []) or []
        task_id     = state.get("task_id", "unknown") or "unknown"

        insights = _extract_json_summary(stdout)
        insights_text = json.dumps(insights, indent=2) if insights else stdout[:2000].strip()

        prompt = f"""Task: {task}

Subtasks completed:
{chr(10).join(f"- {t}" for t in subtasks)}

Research findings (primary source):
{research[:2500] or "(none)"}

{"Execution output:" + chr(10) + insights_text if insights_text else ""}

Artifacts generated: {len(artifacts)} file(s)

{"Sources: " + chr(10).join(f"- {s}" for s in sources[:5]) if sources else ""}

Write a complete, professional Markdown report. Minimum 500 words.
"""

        try:
            response = self.llm.complete(prompt, system_prompt=_SYSTEM_PROMPT)
            report_md = response["content"]
        except Exception as exc:
            logger.error("LLM reporter error: %s", exc)
            report_md = _fallback_md(task, subtasks, research, insights_text)

        # ── Build .docx ────────────────────────────────────────────────────────
        png_artifacts = [a for a in artifacts if a.lower().endswith(".png")]
        total_tokens  = sum(token_usage.values())
        total_time    = sum(timings.values())

        docx_path = _build_docx(
            task_id=task_id,
            task=task,
            report_md=report_md,
            png_paths=png_artifacts,
            code=code,
            sources=sources,
            total_tokens=total_tokens,
            total_time=total_time,
            artifacts_count=len(artifacts),
        )

        pdf_path = build_pdf(
            task_id=task_id,
            task=task,
            report_md=report_md,
            png_paths=png_artifacts,
            sources=sources,
            total_tokens=total_tokens,
            total_time=total_time,
        )

        # Keep markdown version too (used by chat)
        report_md += f"\n\n---\n*Tokens: {total_tokens:,} · Time: {total_time:.1f}s · Artifacts: {len(artifacts)}*\n"
        md_path = _save_text(report_md, task_id)

        return {
            "final_report": report_md,
            "report_path":  str(docx_path) if docx_path else str(md_path),
            "docx_path":    str(docx_path) if docx_path else "",
            "pdf_path":     str(pdf_path) if pdf_path else "",
            "workflow_status": "complete",
            "error_message": "",
        }


# ── DOCX builder ──────────────────────────────────────────────────────────────

def _build_docx(
    task_id: str,
    task: str,
    report_md: str,
    png_paths: list[str],
    code: str,
    sources: list[str],
    total_tokens: int,
    total_time: float,
    artifacts_count: int,
) -> Path | None:
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy
    except ImportError:
        logger.warning("python-docx not installed — skipping .docx generation")
        return None

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # ── Styles ────────────────────────────────────────────────────────────────
    styles = doc.styles

    # Normal text
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Heading 1
    h1 = styles["Heading 1"]
    h1.font.name  = "Calibri"
    h1.font.size  = Pt(16)
    h1.font.bold  = True
    h1.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6e)

    # Heading 2
    h2 = styles["Heading 2"]
    h2.font.name  = "Calibri"
    h2.font.size  = Pt(13)
    h2.font.bold  = True
    h2.font.color.rgb = RGBColor(0x2a, 0x5a, 0x9e)

    # ── Cover page ────────────────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(80)
    run = cover.add_run("Nexus")
    run.font.name  = "Calibri"
    run.font.size  = Pt(32)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x1a, 0x4a, 0xff)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Autonomous Analysis Report")
    r.font.name = "Calibri"
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x4a, 0x6a, 0x9a)

    doc.add_paragraph()

    # Task box
    task_box = doc.add_paragraph()
    task_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = task_box.add_run(f'"{task[:200]}"')
    tr.font.name   = "Calibri"
    tr.font.size   = Pt(11)
    tr.font.italic = True
    tr.font.color.rgb = RGBColor(0x33, 0x33, 0x55)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(f"Tokens: {total_tokens:,}  ·  Time: {total_time:.1f}s  ·  Artifacts: {artifacts_count}")
    mr.font.name  = "Calibri"
    mr.font.size  = Pt(9)
    mr.font.color.rgb = RGBColor(0x88, 0x88, 0x99)

    doc.add_page_break()

    # ── Parse and render Markdown sections ───────────────────────────────────
    _render_markdown(doc, report_md, styles)

    # ── Inline charts ─────────────────────────────────────────────────────────
    valid_pngs = [p for p in png_paths if Path(p).exists()]
    if valid_pngs:
        doc.add_heading("Charts & Visualisations", level=2)
        for png in valid_pngs:
            p = Path(png)
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                cap.add_run().add_picture(str(p), width=Inches(5.5))
            except Exception as exc:
                logger.warning("Could not embed image %s: %s", p.name, exc)
                cap.add_run(f"[Chart: {p.name}]")
            name_para = doc.add_paragraph(p.stem.replace("_", " ").title())
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_para.runs[0].font.size  = Pt(9)
            name_para.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x77)
            doc.add_paragraph()

    # ── Sources ───────────────────────────────────────────────────────────────
    if sources:
        doc.add_heading("Sources", level=2)
        for s in sources[:8]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(s)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1a, 0x4a, 0xff)

    # ── Code appendix ─────────────────────────────────────────────────────────
    if code.strip():
        doc.add_page_break()
        doc.add_heading("Appendix — Generated Code", level=2)
        code_para = doc.add_paragraph()
        code_run  = code_para.add_run(code[:3000])
        code_run.font.name = "Courier New"
        code_run.font.size = Pt(8)
        code_run.font.color.rgb = RGBColor(0x22, 0x44, 0x22)

    # ── Save ──────────────────────────────────────────────────────────────────
    dest = _get_dest_dir(task_id) / f"{task_id}_report.docx"
    try:
        dest.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(dest))
        logger.info("DOCX saved: %s", dest)
        return dest
    except Exception as exc:
        logger.error("Could not save DOCX: %s", exc)
        tmp = Path(tempfile.gettempdir()) / f"{task_id}_report.docx"
        doc.save(str(tmp))
        logger.info("DOCX saved to tmp: %s", tmp)
        return tmp


def _render_markdown(doc, md: str, styles) -> None:
    """Convert Markdown text into Word paragraphs."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, line[2:])
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, re.sub(r"^\d+\.\s", "", line))
        elif line.startswith("```"):
            # Collect code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines[:40]))
                run.font.name = "Courier New"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x22, 0x44, 0x22)
        elif line.startswith("---"):
            # Horizontal rule — skip
            pass
        elif line.strip() == "":
            pass  # blank line
        else:
            p = doc.add_paragraph()
            _add_inline(p, line)

        i += 1


def _add_inline(para, text: str) -> None:
    """Add text with basic bold/italic support."""
    # Split on **bold** and *italic* markers
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            from docx.shared import Pt, RGBColor
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            para.add_run(part)


def _get_dest_dir(task_id: str) -> Path:
    try:
        from config.settings import settings
        return settings.ARTIFACTS_DIR / task_id
    except Exception:
        return Path(tempfile.gettempdir()) / f"amas_{task_id}"


def _save_text(content: str, task_id: str) -> Path:
    path = _get_dest_dir(task_id) / f"{task_id}_report.md"
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content, encoding="utf-8")
    except Exception as exc:
        logger.warning("Could not save MD: %s", exc)
        path = Path(tempfile.gettempdir()) / f"{task_id}_report.md"
        path.write_text(content, encoding="utf-8")
    return path


def _extract_json_summary(stdout: str) -> dict | None:
    if not stdout:
        return None
    for m in reversed(re.findall(r"\{[^{}]*\}", stdout, re.DOTALL)):
        try:
            return json.loads(m)
        except json.JSONDecodeError:
            continue
    return None


def _fallback_md(task: str, subtasks: list, research: str, insights: str) -> str:
    lines = ["# Analysis Report", "", f"## Executive Summary", "", f"Analysis of: {task}", ""]
    if research:
        lines += ["## Key Findings & Insights", "", research[:1200], ""]
    if subtasks:
        lines += ["## Methodology", ""] + [f"- {s}" for s in subtasks] + [""]
    if insights:
        lines += ["## Data Overview", "", f"```\n{insights[:600]}\n```", ""]
    lines += ["## Conclusions & Recommendations", "", "See findings above for detailed conclusions.", ""]
    return "\n".join(lines)


# ── PDF builder ───────────────────────────────────────────────────────────────

def build_pdf(
    task_id: str,
    task: str,
    report_md: str,
    png_paths: list[str],
    sources: list[str],
    total_tokens: int,
    total_time: float,
) -> Path | None:
    """Generate a PDF version of the report using reportlab."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Image,
            HRFlowable, PageBreak, KeepTogether,
        )
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    except ImportError:
        logger.warning("reportlab not installed — skipping PDF")
        return None

    dest = _get_dest_dir(task_id) / f"{task_id}_report.pdf"
    try:
        dest.parent.mkdir(parents=True, exist_ok=True)
    except Exception:
        dest = Path(tempfile.gettempdir()) / f"{task_id}_report.pdf"

    doc = SimpleDocTemplate(
        str(dest),
        pagesize=letter,
        leftMargin=1.1 * inch,
        rightMargin=1.1 * inch,
        topMargin=1.1 * inch,
        bottomMargin=1.1 * inch,
    )

    styles = getSampleStyleSheet()
    BLUE   = colors.HexColor("#1a4aff")
    DARK   = colors.HexColor("#1a2a3a")
    MUTED  = colors.HexColor("#6a7a9a")

    title_style = ParagraphStyle("NTitle", parent=styles["Normal"],
        fontSize=28, textColor=BLUE, spaceAfter=4, fontName="Helvetica-Bold", alignment=TA_CENTER)
    sub_style = ParagraphStyle("NSub", parent=styles["Normal"],
        fontSize=12, textColor=MUTED, spaceAfter=24, alignment=TA_CENTER)
    h1_style = ParagraphStyle("NH1", parent=styles["Normal"],
        fontSize=16, textColor=BLUE, spaceBefore=18, spaceAfter=8, fontName="Helvetica-Bold")
    h2_style = ParagraphStyle("NH2", parent=styles["Normal"],
        fontSize=13, textColor=DARK, spaceBefore=14, spaceAfter=6, fontName="Helvetica-Bold")
    body_style = ParagraphStyle("NBody", parent=styles["Normal"],
        fontSize=10, textColor=DARK, spaceAfter=6, leading=15, alignment=TA_JUSTIFY)
    bullet_style = ParagraphStyle("NBullet", parent=styles["Normal"],
        fontSize=10, textColor=DARK, spaceAfter=4, leftIndent=16, bulletIndent=4, leading=14)
    meta_style = ParagraphStyle("NMeta", parent=styles["Normal"],
        fontSize=8, textColor=MUTED, alignment=TA_CENTER, spaceAfter=4)
    caption_style = ParagraphStyle("NCaption", parent=styles["Normal"],
        fontSize=9, textColor=MUTED, alignment=TA_CENTER, spaceAfter=12)

    story = []

    # Cover
    story.append(Spacer(1, 1.5 * inch))
    story.append(Paragraph("Nexus", title_style))
    story.append(Paragraph("Autonomous Analysis Report", sub_style))
    story.append(HRFlowable(width="60%", thickness=1, color=BLUE, spaceAfter=20, lineCap="round"))
    task_preview = task[:220] + ("…" if len(task) > 220 else "")
    story.append(Paragraph(f'<i>"{task_preview}"</i>', ParagraphStyle("TQ",
        parent=body_style, alignment=TA_CENTER, textColor=DARK, spaceAfter=24)))
    story.append(Paragraph(
        f"Tokens: {total_tokens:,}  ·  Time: {total_time:.1f}s", meta_style))
    story.append(PageBreak())

    # Parse MD into reportlab elements
    for line in report_md.splitlines():
        line = line.rstrip()
        if line.startswith("## "):
            story.append(Paragraph(_escape(line[3:]), h1_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#dde4f0"), spaceAfter=6))
        elif line.startswith("# "):
            story.append(Paragraph(_escape(line[2:]), h1_style))
        elif line.startswith("### "):
            story.append(Paragraph(_escape(line[4:]), h2_style))
        elif line.startswith(("- ", "* ")):
            story.append(Paragraph(f"• {_escape(line[2:])}", bullet_style))
        elif re.match(r"^\d+\.\s", line):
            story.append(Paragraph(_escape(re.sub(r"^\d+\.\s", "", line)), bullet_style))
        elif line.startswith("---"):
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#ccc"), spaceAfter=8))
        elif line.startswith("```") or line.strip() == "":
            pass
        else:
            if line.strip():
                story.append(Paragraph(_escape(line), body_style))

    # Charts
    valid_pngs = [p for p in png_paths if Path(p).exists()]
    if valid_pngs:
        story.append(Spacer(1, 0.3 * inch))
        story.append(Paragraph("Charts & Visualisations", h1_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#dde4f0"), spaceAfter=8))
        for png in valid_pngs:
            p = Path(png)
            try:
                img = Image(str(p), width=5.5 * inch, height=3.5 * inch, kind="proportional")
                story.append(KeepTogether([img, Paragraph(p.stem.replace("_", " ").title(), caption_style)]))
                story.append(Spacer(1, 0.2 * inch))
            except Exception as exc:
                logger.warning("PDF image error %s: %s", p.name, exc)

    # Sources
    if sources:
        story.append(Paragraph("Sources", h1_style))
        for s in sources[:8]:
            story.append(Paragraph(f"• {_escape(s)}", bullet_style))

    try:
        doc.build(story)
        logger.info("PDF saved: %s", dest)
        return dest
    except Exception as exc:
        logger.error("PDF build failed: %s", exc)
        return None


def _escape(text: str) -> str:
    """Escape special XML chars for reportlab."""
    return (text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace('"', "&quot;"))